import pandas as pd
from fastapi.responses import FileResponse
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

#? Create an Excel file from input
def generate_excel(file_data, file_name: str):
    if not file_data or not isinstance(file_data, list):
        return {"status": "failed", "message": "Missing or invalid data"}

    df = pd.DataFrame(file_data)
    file_path = f"{file_name}.xlsx"
    df.to_excel(file_path, index=False)

    return FileResponse(
        path=file_path,
        filename=file_path,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


#? Create a Word file from structured input
def generate_docx(structure, file_name: str):
    if not structure or not isinstance(structure, list):
        return {"status": "failed", "message": "Missing or invalid structure"}

    doc = Document()

    for block in structure:
        block_type = block.get("type")
        if block_type == "heading":
            doc.add_heading(block.get("text", ""), level=block.get("level", 1))
        elif block_type == "paragraph":
            doc.add_paragraph(block.get("text", ""))
        elif block_type == "list":
            for item in block.get("items", []):
                doc.add_paragraph(item, style='ListBullet')

    file_path = f"{file_name}.docx"
    doc.save(file_path)

    return FileResponse(
        path=file_path,
        filename=file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


#? Create a plain text file
def generate_text(content: str, file_name: str):
    if not content:
        return {"status": "failed", "message": "Missing text content"}

    file_path = f"{file_name}.txt"
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)

    return FileResponse(
        path=file_path,
        filename=file_path,
        media_type="text/plain"
    )


def generate_pdf(structure, file_name):
    file_path = f"{file_name}.pdf"
    c = canvas.Canvas(file_path, pagesize=A4)
    width, height = A4
    y = height - 50

    for block in structure:
        if block["type"] == "heading":
            c.setFont("Helvetica-Bold", 16)
            c.drawString(50, y, block["text"])
            y -= 30
        elif block["type"] == "paragraph":
            c.setFont("Helvetica", 12)
            for line in block["text"].split('\n'):
                c.drawString(50, y, line)
                y -= 20
        elif block["type"] == "list":
            c.setFont("Helvetica", 12)
            for item in block["items"]:
                c.drawString(70, y, f"• {item}")
                y -= 20

        y -= 10  # Add spacing between blocks
        if y < 100:
            c.showPage()
            y = height - 50

    c.save()

    return FileResponse(
        path=file_path,
        filename=file_path,
        media_type="application/pdf"
    )